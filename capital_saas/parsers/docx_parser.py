from pathlib import Path

from docx import Document

from parsers.base_parser import BaseParser


class DocxParser(BaseParser):
    parser_type = "docx"

    def parse(self, file_path: Path) -> dict:
        if file_path.suffix.lower() == ".doc":
            raise ValueError("旧版 .doc 暂不支持，请另存为 .docx 后重新上传")
        document = Document(file_path)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        tables = []
        for table in document.tables[:20]:
            tables.append([[cell.text.strip() for cell in row.cells] for row in table.rows[:50]])
        text = "\n".join(paragraphs)
        return {
            "parser_type": self.parser_type,
            "status": "success",
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs[:100],
            "tables": tables,
            "text_summary": text[:5000],
        }
