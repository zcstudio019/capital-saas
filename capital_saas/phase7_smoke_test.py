"""Phase 7：资料解析、完整性、尽调、补全建议与申请材料包验收。"""
import io
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import DictionaryObject, NameObject, DecodedStreamObject

from db.database import SessionLocal
from db.models import (DocumentParseTask, DueDiligenceReport, FinancingApplicationPackage,
    FollowTask, Lead, LeadTag, UploadedDocument)
from main import app

PAYLOAD = {"company_name":"Phase7尽调测试企业有限公司","contact_name":"资料负责人","phone":"13400134000",
"wechat_id":"phase7_dd","city":"上海","industry":"制造业","years":"6","employee_count":"55",
"annual_revenue":"10000000","net_profit":"800000","monthly_cashflow":"250000","debt_total":"5000000",
"short_debt":"3500000","receivable_days":"100","funding_need":"5000000","funding_purpose":"订单周转",
"has_collateral":"true","tax_status":"true","credit_status":"true","knows_cashflow":"true",
"has_budget":"false","leverage_attitude":"适中","asset_efficiency":"中","fund_usage_plan":"true"}


def xlsx_bytes():
    wb = Workbook(); ws = wb.active; ws.title = "利润表"
    ws.append(["项目", "本年累计"]); ws.append(["营业收入", 18000000]); ws.append(["净利润", 1500000])
    ws.append(["总负债", 6200000]); ws.append(["短期借款", 3200000]); ws.append(["现金流", 480000])
    output = io.BytesIO(); wb.save(output); return output.getvalue()


def docx_bytes():
    doc = Document(); doc.add_heading("企业工商资料", 1); doc.add_paragraph("统一社会信用代码 91310000PHASE7")
    table = doc.add_table(rows=1, cols=2); table.rows[0].cells[0].text = "企业名称"; table.rows[0].cells[1].text = PAYLOAD["company_name"]
    output = io.BytesIO(); doc.save(output); return output.getvalue()


def pdf_bytes():
    writer = PdfWriter(); page = writer.add_blank_page(width=595, height=842)
    font = DictionaryObject({NameObject("/Type"): NameObject("/Font"), NameObject("/Subtype"): NameObject("/Type1"), NameObject("/BaseFont"): NameObject("/Helvetica")})
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})})
    stream = DecodedStreamObject(); stream.set_data(b"BT /F1 12 Tf 72 720 Td (Bank Statement Phase7 Cash Flow) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(stream)
    output = io.BytesIO(); writer.write(output); return output.getvalue()


def login(client):
    assert client.post("/login", data={"username":"admin","password":"admin123","next_url":"/admin"}, follow_redirects=False).status_code == 303


def upload(client, lead_id, category, filename, content, mime):
    return client.post(f"/admin/leads/{lead_id}/document-center/upload",
        data={"document_category":category,"note":"Phase7验收资料"},
        files=[("uploads",(filename,io.BytesIO(content),mime))], follow_redirects=False)


def run():
    with TestClient(app) as client:
        submit = client.post("/assessment/submit", data=PAYLOAD, follow_redirects=False)
        assessment_id = int(submit.headers["location"].rsplit("/",1)[-1]); login(client)
        with SessionLocal() as db:
            lead = db.query(Lead).filter(Lead.assessment_id == assessment_id).one(); lead_id = lead.id
        assert client.get(f"/admin/leads/{lead_id}/document-center").status_code == 200
        excel = xlsx_bytes()
        assert upload(client,lead_id,"财务报表","2025财务报表.xlsx",excel,"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet").status_code == 303
        assert upload(client,lead_id,"营业执照/工商资料","营业执照.docx",docx_bytes(),"application/vnd.openxmlformats-officedocument.wordprocessingml.document").status_code == 303
        assert upload(client,lead_id,"银行流水","银行流水.pdf",pdf_bytes(),"application/pdf").status_code == 303
        assert upload(client,lead_id,"纳税资料","完税证明.png",b"\x89PNG\r\n\x1a\nphase7","image/png").status_code == 303
        duplicate = upload(client,lead_id,"财务报表","财务报表副本.xlsx",excel,"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        assert "duplicate=1" in duplicate.headers["location"]

        with SessionLocal() as db:
            docs = db.query(UploadedDocument).filter(UploadedDocument.lead_id == lead_id).all()
            assert len(docs) == 5 and db.query(DocumentParseTask).filter(DocumentParseTask.lead_id == lead_id).count() == 5
            excel_doc = next(x for x in docs if x.file_name == "2025财务报表.xlsx")
            word_doc = next(x for x in docs if x.file_name == "营业执照.docx")
            pdf_doc = next(x for x in docs if x.file_name == "银行流水.pdf")
            image_doc = next(x for x in docs if x.file_name == "完税证明.png")
            assert json.loads(excel_doc.parsed_json)["financial_fields"]["annual_revenue"] == 18000000
            assert "企业工商资料" in json.loads(word_doc.parsed_json)["text_summary"]
            assert "Bank Statement" in json.loads(pdf_doc.parsed_json)["text_summary"]
            assert json.loads(image_doc.parsed_json)["ocr_status"] == "pending_ocr"
            excel_id, pdf_id = excel_doc.id, pdf_doc.id
        assert client.post(f"/admin/documents/{excel_id}/reparse", follow_redirects=False).status_code == 303
        assert client.post(f"/admin/documents/{pdf_id}/verify", data={"verify_status":"verified"}, follow_redirects=False).status_code == 303

        assert client.post(f"/admin/leads/{lead_id}/due-diligence/generate", follow_redirects=False).status_code == 303
        dd_page = client.get(f"/admin/leads/{lead_id}/due-diligence")
        assert dd_page.status_code == 200 and "资料完整度" in dd_page.text
        review = client.get(f"/admin/leads/{lead_id}/autofill-review")
        assert review.status_code == 200 and "annual_revenue" in review.text
        assert client.post(f"/admin/leads/{lead_id}/autofill/apply", data={"fields":["annual_revenue","net_profit"]}, follow_redirects=False).status_code == 303

        package = client.post(f"/admin/leads/{lead_id}/application-package/create", data={
            "package_name":"Phase7融资申请包","target_product_code":"1999_structure_plan",
            "target_bank_product_id":"0","document_ids":[str(excel_id),str(pdf_id)],
            "advisor_note":"先预审后申请"}, follow_redirects=False)
        assert package.status_code == 303
        with SessionLocal() as db:
            lead = db.get(Lead, lead_id); assert lead.assessment.annual_revenue == 18000000
            dd = db.query(DueDiligenceReport).filter(DueDiligenceReport.lead_id == lead_id).one()
            pkg = db.query(FinancingApplicationPackage).filter(FinancingApplicationPackage.lead_id == lead_id).one()
            assert dd.completeness_score >= 0
            assert db.query(FollowTask).filter(FollowTask.lead_id == lead_id, FollowTask.task_type == "collect_documents").count() >= 1
            tags = {link.tag.name for link in db.query(LeadTag).filter(LeadTag.lead_id == lead_id).all()}
            assert "资料不完整" in tags and "需人工尽调" in tags
            package_id = pkg.id
        csv_response = client.get(f"/admin/application-packages/{package_id}/checklist.csv")
        assert csv_response.status_code == 200 and csv_response.content.startswith(b"\xef\xbb\xbf")
        assert client.get(f"/admin/leads/{lead_id}/application-package").status_code == 200
        assert client.post("/api/events/document-request-script-copied", data={"lead_id":lead_id}).status_code == 200
        print({"assessment_id":assessment_id,"lead_id":lead_id,"documents":5,"package_id":package_id,"dd_score":dd.completeness_score})
    print("PHASE7_DOCUMENT_DUE_DILIGENCE_OK")


if __name__ == "__main__": run()
